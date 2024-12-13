from flask import Flask, render_template, request
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from docx.shared import Pt
import os
import sys
from docx.shared import Inches
import base64


### Working directory ###
if getattr(sys, 'frozen', False):
    workdir = os.path.dirname(sys.executable)
else:
    workdir = os.path.dirname(os.path.abspath(__file__))


inventoryvico = Flask(__name__)

@inventoryvico.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        technitian_name = request.form['technitian_name']
        input1 = request.form['input1']
        input2 = request.form['input2']
        input3 = request.form['input3']
        input4 = request.form['input4']
        input5 = request.form['input5']
        input6 = request.form['input6']
        input7 = request.form['input7']
        input8 = request.form['input8']
        input9 = request.form['input9']
        input10 = request.form['input10']
        input11 = request.form['input11']
        input12 = request.form['input12']
        input13 = request.form['input13']
        input14 = request.form['input14']
        input15 = request.form['input15']
        input16 = request.form['input16']
        input17 = request.form['input17']
        input18 = request.form['input18']
        input19 = request.form['input19']
        input20 = request.form['input20']
        input21 = request.form['input21']
        input22 = request.form['input22']
        input23 = request.form['input23']
        input24 = request.form['input24']
        input25 = request.form['input25']
        input26 = request.form['input26']
        input27 = request.form['input27']
        input28 = request.form['input28']
        input29 = request.form['input29']
        input30 = request.form['input30']
        input31 = request.form['input31']
        input32 = request.form['input32']
        tech_remark = request.form['tech_remark']

        # Create a new document
        document = Document()

        # Set RTL direction for the document
        document.core_properties.language = 'he-IL'  # Set to Hebrew (Israel)
        document.styles['Normal'].font.name = 'Tahoma'  # Replace with a Hebrew font name
        document.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add a heading
        document.add_heading(f'טופס ספירת מלאי')

        # Modify the font size of the heading
        heading_paragraph = document.paragraphs[0]  # Assuming the heading is the first paragraph
        heading_paragraph.style.font.size = Pt(18)  # Replace 16 with the desired font size in points

        # Set RTL direction for the document (corrected)
        document.settings.element.rtl = True

        # Add a header with the current date and the name
        header = document.sections[0].header

        # Add logo to the header
        logo_path = os.path.join(workdir, 'pics', 'logo.png')
        logo_width = Inches(1.9)
        logo_height = Inches(0.5)
        logo_paragraph = header.add_paragraph()
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, width=logo_width, height=logo_height)
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add the current date to the header
        date_paragraph = header.add_paragraph()
        date_paragraph.text = datetime.now().strftime('%d.%m.%Y') + ' - ' + technitian_name# Set the text to the current date
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add the inputs to the document
        table = document.add_table(rows=32, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        table.style = 'Table Grid'  # add table style

        for row in table.rows:
            for cell in row.cells:
                cell = table.cell(0, 0)
                cell_text = "Text in cell"
                paragraph = cell.add_paragraph(cell_text)
                paragraph.style = document.styles['Normal']
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph.runs[0].font.size = Pt(13)


        table.cell(0, 0).text = input1
        table.cell(0, 1).text = 'קופה 5200 שחורה חדש'
        table.cell(1, 0).text = input2
        table.cell(1, 1).text = 'קופה 5200 שחורה שכירות'
        table.cell(2, 0).text = input3
        table.cell(2, 1).text = 'קופה 5200 לבנה חדש'
        table.cell(3, 0).text = input4
        table.cell(3, 1).text = 'קופה 5200 לבנה שכירות'
        table.cell(4, 0).text = input5
        table.cell(4, 1).text = 'מגירות שכירות + רכישה'
        table.cell(5, 0).text = input6
        table.cell(5, 1).text = 'P400 חדש'
        table.cell(6, 0).text = input7
        table.cell(6, 1).text = 'P400 שכירות'
        table.cell(7, 0).text = input8
        table.cell(7, 1).text = 'VX680'
        table.cell(8, 0).text = input9
        table.cell(8, 1).text = 'VX820'
        table.cell(9, 0).text = input10
        table.cell(9, 1).text = 'VX520 GPRS'
        table.cell(10, 0).text = input11
        table.cell(10, 1).text = 'VX520 ETH'
        table.cell(11, 0).text = input12
        table.cell(11, 1).text = 'VX520 DEBIT'
        table.cell(12, 0).text = input13 
        table.cell(12, 1).text = 'VX805'
        table.cell(13, 0).text = input14
        table.cell(13, 1).text = ' סוויץ רשת'
        table.cell(14, 0).text = input15
        table.cell(14, 1).text = 'נתב סלולארי דור 4 '
        table.cell(15, 0).text = input16
        table.cell(15, 1).text = 'אקדח ברקוד ידני'
        table.cell(16, 0).text = input17
        table.cell(16, 1).text = 'קופה 6200'
        table.cell(17, 0).text = input18
        table.cell(17, 1).text = 'CM5 קופה'
        table.cell(18, 0).text = input19
        table.cell(18, 1).text = 'CM5 אשראית'
        table.cell(19, 0).text = input20
        table.cell(19, 1).text = 'T650 קופה'
        table.cell(20, 0).text = input21
        table.cell(20, 1).text = 'T650 אשראית'
        table.cell(21, 0).text = input22
        table.cell(21, 1).text = 'T650C אשראית'
        table.cell(22, 0).text = input23
        table.cell(22, 1).text = 'T650 סיבוס'
        table.cell(23, 0).text = input24
        table.cell(23, 1).text = 'T650C סיבוס'
        table.cell(24, 0).text = input25
        table.cell(24, 1).text = 'CM5 קופה שכירות'
        table.cell(25, 0).text = input26
        table.cell(25, 1).text = 'CM5 אשראית שכירות'
        table.cell(26, 0).text = input27
        table.cell(26, 1).text = '650T אשראית שכירות'
        table.cell(27, 0).text = input28
        table.cell(27, 1).text = 'AI T650 קופה/אשראית שכירות'
        table.cell(28, 0).text = input29
        table.cell(28, 1).text = 'C650T שכירות'
        table.cell(29, 0).text = input30
        table.cell(29, 1).text = 'נייר קטן'
        table.cell(30, 0).text = input31
        table.cell(30, 1).text = 'נייר גדול'
        table.cell(31, 0).text = input32
        table.cell(31, 1).text = 'נייר ענק גזית'


        # # Add a blank paragraph outside the table
        document.add_paragraph()
        # document.add_paragraph()
        # document.add_paragraph()

        # Add a paragraph with the text from tech_remark
        default_text = """הערות טכנאי """
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        run = paragraph.add_run(default_text + " : ")
        run.font.size = Pt(12)
        run.bold = True
        paragraph.add_run(" " * 6 + tech_remark).font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        ##############################  SIGNATURE ########################
        # Create a title paragraph for the signature
        title = document.add_paragraph(style='Normal')
        title.add_run('\n\n\n\n:חתימת הטכנאי')    
        file_path = os.path.join(workdir, 'files', 'signature.png')
        # Save signature as image file
        with open(file_path, 'wb') as f:
            signature_data = request.form['signature_data']
            signature_data = signature_data.replace('data:image/png;base64,', '')
            f.write(base64.b64decode(signature_data))

        # Add signature image to document
        document.add_picture(file_path, width=Inches(2.0))


        # Save the document
        filename = technitian_name + '_' + datetime.now().strftime('%d-%m-%Y') + '.docx'
        document.save(workdir + '//files//' + filename)
        filename_without_path = technitian_name + '_' + datetime.now().strftime('%d-%m-%Y') + '.docx'

        
        #########################################   Send email #############################################

        import email, smtplib, ssl
        from email import encoders
        from email.mime.base import MIMEBase
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        from email.message import EmailMessage


        subject = f"טופס ספירת מלאי רכב ({technitian_name})"
        body = "VICO Field services team"
        sender_email = "vrics@verifone.co.il"
        receiver_email = ["dimitry_d1@VERIFONE.com"]
        # receiver_email = ["crooper22@gmail.com"]

        # Create a multipart message and set headers
        message = MIMEMultipart()
        message["From"] = "VICO Field services team<{}>".format(sender_email)
        message["To"] = ', '.join(receiver_email)
        message["Subject"] = subject
        message.attach(MIMEText(body, "plain"))

        filename = os.path.join(workdir, "files", filename)
        filename_without_path = technitian_name + '_' + datetime.now().strftime('%d-%m-%Y') + '.docx'

        # Open the file in binary mode and read its contents
        with open(filename, "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
            encoders.encode_base64(part)

        # Encode the filename with UTF-8
        filename_encoded = filename_without_path.encode("utf-8")

        # Set the Content-Disposition header with the encoded filename
        part.add_header(
            "Content-Disposition",
            "attachment",
            filename=("utf-8", "", filename_encoded.decode("utf-8"))
        )

        message.attach(part)
        text = message.as_string()

        with smtplib.SMTP("verifonecp.agas.co.il", port=587) as server:
            server.starttls()
            server.login(sender_email, "(K$6&Mu($^cB")
            server.sendmail(sender_email, receiver_email, text)


        
        ### remove temporary file ###
        if os.path.exists(filename):
            os.remove(filename)
##################################################################################

        return '''<html>
        <head>
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body>
            <div style="
                                display: flex;
                                justify-content: center;
                                align-items: center;
                                height: 100%;
                                font-size: 32px;
                                font-weight: bold;
                            ">
            File sent by email to BO VICO {}
            </div>
        </body>
        </html>
        '''.format(filename_without_path)
    
    return render_template('index.html')

if __name__ == '__main__':
    inventoryvico.run(host='0.0.0.0', port=5003)
